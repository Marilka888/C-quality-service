"""
Batch parse 138 folders (each folder contains 2 files: TZ + PMI, pdf/docx, arbitrary filenames)
into JSON packages: data/packages/<folder>.json

Key idea: classify document type by title/first pages text (not by filename).

Usage:
python scripts/batch_parse_packages.py --raw_root "D:/data/raw_packages" --out_dir "./data/packages" --min_chunks 20

Outputs:
- ./data/packages/*.json
- ./data/batch_parse_report.json (full per-folder diagnostics)
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional

import docx
import fitz  # PyMuPDF


# -----------------------------
# 1) Title-based classification patterns
# -----------------------------

TZ_STRONG = [
    r"техническ(ое|ого)\s+задан",          # "Техническое задание"
    r"\bтз\b",                             # "ТЗ" as token (often in doc code)
    r"гост\s*19\.201",                     # GOST for TZ
    r"основан(ие|ия)\s+для\s+разработк",   # typical TZ section
    r"требован(ия|ие)\s+к\s+програм",      # typical TZ section
    r"порядок\s+контроля\s+и\s+приемки",   # typical TZ section
]

PMI_STRONG = [
    r"программ(а|ы)\s+и\s+метод(ика|ики)\s+испытан",  # "Программа и методика испытаний"
    r"гост\s*19\.301",                                 # GOST for PMI
    r"\bобъект\s+испытан",                             # typical PMI section
    r"\bцель\s+испытан",                               # typical PMI section
    r"\bметод(ы|ика)\s+испытан",                       # typical PMI section
]

TZ_RE = re.compile("|".join(TZ_STRONG), flags=re.IGNORECASE | re.UNICODE)
PMI_RE = re.compile("|".join(PMI_STRONG), flags=re.IGNORECASE | re.UNICODE)


def head_text_pdf(path: Path, pages: int = 2, max_chars: int = 15000) -> str:
    doc = fitz.open(str(path))
    out = []
    for i in range(min(pages, len(doc))):
        out.append(doc[i].get_text("text") or "")
    return ("\n".join(out))[:max_chars]


def head_text_docx(path: Path, paragraphs: int = 60, max_chars: int = 15000) -> str:
    d = docx.Document(str(path))
    out = []
    for p in d.paragraphs[:paragraphs]:
        t = (p.text or "").strip()
        if t:
            out.append(t)
    return ("\n".join(out))[:max_chars]


def head_text(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".pdf":
        return head_text_pdf(path)
    if suf == ".docx":
        return head_text_docx(path)
    return ""


def score_types(text: str) -> Tuple[int, int, Dict[str, int]]:
    """
    Returns tz_score, pmi_score + debug counts
    """
    t = text.lower()

    tz_hits = len(re.findall(TZ_RE, t))
    pmi_hits = len(re.findall(PMI_RE, t))

    # Big bonuses for exact title phrases (high precision)
    if re.search(r"техническ(ое|ого)\s+задан", t):
        tz_hits += 5
    if re.search(r"программ(а|ы)\s+и\s+метод(ика|ики)\s+испытан", t):
        pmi_hits += 5

    dbg = {
        "tz_hits": tz_hits,
        "pmi_hits": pmi_hits,
        "has_tz_title": 1 if re.search(r"техническ(ое|ого)\s+задан", t) else 0,
        "has_pmi_title": 1 if re.search(r"программ(а|ы)\s+и\s+метод(ика|ики)\s+испытан", t) else 0,
        "has_gost_19201": 1 if re.search(r"гост\s*19\.201", t) else 0,
        "has_gost_19301": 1 if re.search(r"гост\s*19\.301", t) else 0,
    }
    return tz_hits, pmi_hits, dbg


def classify_file_by_title(fp: Path, min_conf: int = 2) -> Tuple[str, int, int, Dict[str, int]]:
    """
    returns (dtype, tz_score, pmi_score, dbg)
    dtype in TZ/PMI/OTHER
    """
    h = head_text(fp)
    tz, pmi, dbg = score_types(h)

    if tz > pmi and tz >= min_conf:
        return "TZ", tz, pmi, dbg
    if pmi > tz and pmi >= min_conf:
        return "PMI", tz, pmi, dbg
    return "OTHER", tz, pmi, dbg


# -----------------------------
# 2) Chunk extraction (full doc)
# -----------------------------

def chunks_from_docx(path: Path) -> List[Dict]:
    d = docx.Document(str(path))
    chunks: List[Dict] = []
    idx = 1
    for p in d.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        chunks.append({
            "chunk_id": f"{path.stem}_{idx:05d}",
            "section_path": None,
            "text": t,
            "order": idx,
            "page": None,
        })
        idx += 1
    return chunks


def chunks_from_pdf(path: Path) -> List[Dict]:
    doc = fitz.open(str(path))
    chunks: List[Dict] = []
    idx = 1
    for page_i in range(len(doc)):
        text = doc[page_i].get_text("text") or ""
        # paragraph-like splits
        parts = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        for p in parts:
            chunks.append({
                "chunk_id": f"{path.stem}_{idx:05d}",
                "section_path": None,
                "text": p,
                "order": idx,
                "page": page_i + 1,
            })
            idx += 1
    return chunks


def parse_file_to_chunks(fp: Path) -> List[Dict]:
    suf = fp.suffix.lower()
    if suf == ".docx":
        return chunks_from_docx(fp)
    if suf == ".pdf":
        return chunks_from_pdf(fp)
    return []


# -----------------------------
# 3) Folder-level selection (2 files)
# -----------------------------

def pick_tz_pmi(files: List[Path], min_conf: int) -> Tuple[Optional[Path], Optional[Path], List[Dict]]:
    """
    Classify each file by title text and pick best TZ and PMI by score.
    Returns tz_fp, pmi_fp, debug list.
    """
    scored = []
    for fp in files:
        dtype, tz_s, pmi_s, dbg = classify_file_by_title(fp, min_conf=min_conf)
        scored.append({
            "file": fp.name,
            "dtype": dtype,
            "tz_score": tz_s,
            "pmi_score": pmi_s,
            "dbg": dbg
        })

    # candidates
    tz_best = None
    pmi_best = None

    # best by tz_score / pmi_score regardless of dtype (fallback)
    tz_best_by_score = max(files, key=lambda f: next(x["tz_score"] for x in scored if x["file"] == f.name), default=None)
    pmi_best_by_score = max(files, key=lambda f: next(x["pmi_score"] for x in scored if x["file"] == f.name), default=None)

    # best by explicit dtype
    tz_cands = [f for f in files if any(x["file"] == f.name and x["dtype"] == "TZ" for x in scored)]
    pmi_cands = [f for f in files if any(x["file"] == f.name and x["dtype"] == "PMI" for x in scored)]

    if tz_cands:
        tz_best = max(tz_cands, key=lambda f: next(x["tz_score"] for x in scored if x["file"] == f.name))
    if pmi_cands:
        pmi_best = max(pmi_cands, key=lambda f: next(x["pmi_score"] for x in scored if x["file"] == f.name))

    # if one missing — fallback to by-score
    if tz_best is None:
        tz_best = tz_best_by_score
    if pmi_best is None:
        pmi_best = pmi_best_by_score

    # if they are the same file (rare but possible if all OTHER) -> unresolved
    if tz_best is not None and pmi_best is not None and tz_best.resolve() == pmi_best.resolve():
        return None, None, scored

    return tz_best, pmi_best, scored


# -----------------------------
# 4) Main batch
# -----------------------------

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--raw_root", required=True, help="Root folder with package folders")
    ap.add_argument("--out_dir", required=True, help="Output folder for data/packages/*.json")
    ap.add_argument("--min_chunks", type=int, default=20, help="Skip docs with too few chunks")
    ap.add_argument("--min_conf", type=int, default=2, help="Min confidence score for title-based type")
    args = ap.parse_args()

    raw_root = Path(args.raw_root)
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    report = {
        "total_folders": 0,
        "ok": 0,
        "ambiguous": 0,
        "skipped_small": 0,
        "errors": 0,
        "items": []
    }

    for folder in sorted([p for p in raw_root.iterdir() if p.is_dir()]):
        report["total_folders"] += 1

        files = [f for f in folder.iterdir()
                 if f.is_file() and f.suffix.lower() in (".pdf", ".docx")]

        # If there are more than 2 files, we still handle it (pick best by title)
        if len(files) < 2:
            report["ambiguous"] += 1
            report["items"].append({"folder": folder.name, "status": "AMBIGUOUS_TOO_FEW_FILES", "files": [f.name for f in files]})
            continue

        try:
            tz_fp, pmi_fp, dbg = pick_tz_pmi(files, min_conf=args.min_conf)

            if tz_fp is None or pmi_fp is None:
                report["ambiguous"] += 1
                report["items"].append({
                    "folder": folder.name,
                    "status": "AMBIGUOUS_CLASSIFY",
                    "files": [f.name for f in files],
                    "scoring": dbg
                })
                continue

            tz_chunks = parse_file_to_chunks(tz_fp)
            pmi_chunks = parse_file_to_chunks(pmi_fp)

            if len(tz_chunks) < args.min_chunks or len(pmi_chunks) < args.min_chunks:
                report["skipped_small"] += 1
                report["items"].append({
                    "folder": folder.name,
                    "status": "SKIP_SMALL",
                    "tz_file": tz_fp.name,
                    "pmi_file": pmi_fp.name,
                    "tz_chunks": len(tz_chunks),
                    "pmi_chunks": len(pmi_chunks),
                    "scoring": dbg
                })
                continue

            pkg = {
                "package_id": folder.name,
                "documents": [
                    {"doc_id": tz_fp.stem, "doc_type": "TZ", "chunks": tz_chunks},
                    {"doc_id": pmi_fp.stem, "doc_type": "PMI", "chunks": pmi_chunks},
                ],
            }

            out_path = out_dir / f"{folder.name}.json"
            out_path.write_text(json.dumps(pkg, ensure_ascii=False, indent=2), encoding="utf-8")

            report["ok"] += 1
            report["items"].append({
                "folder": folder.name,
                "status": "OK",
                "tz_file": tz_fp.name,
                "pmi_file": pmi_fp.name,
                "tz_chunks": len(tz_chunks),
                "pmi_chunks": len(pmi_chunks),
                "out": out_path.name,
                "scoring": dbg
            })

        except Exception as e:
            report["errors"] += 1
            report["items"].append({
                "folder": folder.name,
                "status": "ERROR",
                "error": str(e),
                "files": [f.name for f in files]
            })

    report_path = out_dir.parent / "batch_parse_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print("DONE")
    print(json.dumps({k: report[k] for k in report if k != "items"}, ensure_ascii=False, indent=2))
    print("Full report:", str(report_path))


if __name__ == "__main__":
    main()
